import os
import re
import json
import hashlib
import unicodedata
from pathlib import Path
from typing import Dict, Any

from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document

import re
from datetime import datetime
from typing import Iterable, List, Optional

def normalize_one_line(s: Optional[str]) -> str:
    if s is None:
        return ""
    s = str(s)
    s = re.sub(r"\s+", " ", s)  
    return s.strip()

def line(label: str, value: Optional[str]) -> Optional[str]:
    v = normalize_one_line(value)
    if not v:
        return None
    return f"- {label}: {v}"

def to_str_list(x) -> List[str]:
    if x is None:
        return []
    if isinstance(x, str):
        parts = re.split(r"[;\n•,]+", x)
        return [normalize_one_line(p) for p in parts if normalize_one_line(p)]
    if isinstance(x, Iterable):
        out = []
        for it in x:
            v = normalize_one_line(it)
            if v:
                out.append(v)
        return out
    v = normalize_one_line(str(x))
    return [v] if v else []

def try_format_date_yyyy_mm_dd(value: Optional[str]) -> str:
    v = normalize_one_line(value)
    if not v:
        return ""
    try:
        return datetime.fromisoformat(v[:10]).strftime("%Y-%m-%d")
    except Exception:
        pass

    m = re.match(r"^(\d{1,2})[/-](\d{1,2})[/-](\d{4})$", v)
    if m:
        d, mth, y = m.groups()
        try:
            return datetime(int(y), int(mth), int(d)).strftime("%Y-%m-%d")
        except Exception:
            return v

    return v

# --- Générateur de contenu ------------------------------------------------

def make_contenu_from_candidat(c) -> str:
    ip   = getattr(c, "informations_personnelles", None)
    exps = getattr(c, "experiences_professionnelles", []) or []
    frms = getattr(c, "formation", []) or []
    lngs = getattr(c, "langues", []) or []
    cert = getattr(c, "certifications", []) or []
    hbs  = getattr(c, "passe_temps", []) or []
    comp = getattr(c, "competences", None)

    header_name = " ".join(filter(None, [
        normalize_one_line(getattr(ip, "prenom", "") if ip else ""),
        normalize_one_line(getattr(ip, "nom", "") if ip else "")
    ])).strip()
    lines = [("CV — " + header_name).strip(" —")]

    ident = list(filter(None, [
        line("Sexe", getattr(ip, "sexe", None) if ip else None),
        line("Date de naissance", try_format_date_yyyy_mm_dd(getattr(ip, "date_de_naissance", None) if ip else None)),
        line("Adresse", getattr(ip, "adresse", None) if ip else None),
        line("Email", getattr(ip, "email", None) if ip else None),
        line("Téléphone", getattr(ip, "numero_telephone", None) if ip else None),
    ]))
    if ident:
        lines += ["", "IDENTITÉ", *ident]

    exp_lines: List[str] = []
    for e in exps:
        t = normalize_one_line(getattr(e, "titre_du_poste", ""))
        c_ = normalize_one_line(getattr(e, "nom_de_l_entreprise", ""))
        d = normalize_one_line(getattr(e, "dates_de_debut_et_de_fin", ""))

        head_parts = [p for p in [t, c_] if p]
        head = " — ".join(head_parts)
        if d:
            head = f"{head} ({d})" if head else d

        missions = getattr(e, "missions_realisees", [])
        missions = missions if isinstance(missions, list) else to_str_list(missions)

        if head:
            exp_lines.append(f"- {head}")
            exp_lines += [f"   • {m}" for m in missions if m]
    if exp_lines:
        lines += ["", "EXPÉRIENCES", *exp_lines]

    frm_lines: List[str] = []
    for f in frms:
        diplome = normalize_one_line(getattr(f, "diplome", ""))
        etab   = normalize_one_line(getattr(f, "etablissement", ""))
        dates  = normalize_one_line(getattr(f, "dates_de_debut_et_de_fin", ""))
        spec   = normalize_one_line(getattr(f, "specialite", ""))

        parts = [p for p in [diplome, etab] if p]
        line1 = " — ".join(parts) if parts else ""
        if dates:
            line1 = f"{line1} ({dates})" if line1 else dates
        if spec:
            line1 = f"{line1} — Spécialité: {spec}" if line1 else f"Spécialité: {spec}"

        if line1:
            frm_lines.append(f"- {line1}")
    if frm_lines:
        lines += ["", "FORMATION", *frm_lines]

    lng_lines: List[str] = []
    for l in lngs:
        a = normalize_one_line(getattr(l, "langue_maitrisee", ""))
        n = normalize_one_line(getattr(l, "niveau_de_maitrise", ""))
        item = f"- {a}: {n}".strip(": ").strip()
        if item != "-":
            if a or n:
                lng_lines.append(item)
    if lng_lines:
        lines += ["", "LANGUES", *lng_lines]

    cert_lines: List[str] = []
    for cc in cert:
        t = normalize_one_line(getattr(cc, "titre", ""))
        d = normalize_one_line(getattr(cc, "dates_d_obtention", ""))
        item = t if t else ""
        if d:
            item = f"{item} ({d})" if item else d
        if item:
            cert_lines.append(f"- {item}")
    if cert_lines:
        lines += ["", "CERTIFICATIONS", *cert_lines]

    tech = []
    soft = []
    if comp is not None:
        tech = [normalize_one_line(x) for x in (getattr(comp, "competences_techniques", []) or []) if normalize_one_line(x)]
        soft = [normalize_one_line(x) for x in (getattr(comp, "competences_comportementales_soft_skills", []) or []) if normalize_one_line(x)]
    if tech or soft:
        lines += ["", "COMPÉTENCES"]
        if tech:
            lines += ["- Techniques: " + ", ".join(tech)]
        if soft:
            lines += ["- Comportementales: " + ", ".join(soft)]

    hb_lines = []
    for h in hbs:
        nom = normalize_one_line(getattr(h, "nom", ""))
        if nom:
            hb_lines.append(f"- {nom}")
    if hb_lines:
        lines += ["", "CENTRES D'INTÉRÊT", *hb_lines]

    txt = "\n".join(lines)
    txt = re.sub(r"\n{3,}", "\n\n", txt).strip()
    return txt

# --- Extraction de texte brut ----------------------------------------------

def extract_text_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("  ".join(cell.text for cell in row.cells))
    return "\n".join(parts)

def extract_text_pdf(file_path: str) -> str:
    return extract_pdf_text(file_path)

DEFAULT_CONFIG = {
    "normalize_unicode": True,
    "strip_zero_width": True,
    "replace_nbsp": True,
    "smart_quotes_to_ascii": False,  
    "deduplicate_spaces": True,
    "collapse_blank_lines": True,
    "fix_hyphenation": True,   
    "join_broken_lines": True,     
    "remove_bullets": True,          
    "remove_headers_footers": True, 
    "trim_lines": True,
    "max_length": None,          
    "redact_emails": False,
    "redact_phones": False,
}

ZERO_WIDTH_PATTERN = re.compile(r"[\u200B-\u200F\u202A-\u202E\u2060-\u206F]")
NBSP_PATTERN = re.compile(r"\u00A0")
BULLET_PATTERN = re.compile(r"^\s*([•·\-–—\*]\s+)", re.MULTILINE)
HEADER_FOOTER_CANDIDATE = re.compile(
    r"^(page\s*\d+(/\d+)?|confidentiel|draft|footer|header)\b.*$",
    re.IGNORECASE | re.MULTILINE
)
EMAIL_PATTERN = re.compile(r"\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[A-Za-z]{2,}\b")
PHONE_PATTERN = re.compile(r"\b(?:\+?\d{1,3}[\s\-\.]?)?(?:\(?\d{2,4}\)?[\s\-\.]?)?\d{3,4}[\s\-\.]?\d{3,4}\b")

def smart_quotes_to_ascii(s: str) -> str:
    pairs = {
        """: '"', """: '"', "„": '"', "‟": '"', "«": '"', "»": '"',
        "'": "'", "'": "'", "‚": "'", "‛": "'",
        "–": "-", "—": "-", "−": "-", "…": "...",
    }
    for k, v in pairs.items():
        s = s.replace(k, v)
    return s

def fix_hyphenation(text: str) -> str:
    return re.sub(r"(\w)-\n(\w)", r"\1\2", text)

def join_broken_lines(text: str) -> str:
    return re.sub(r"(?<![\.!\?:;\)\]])\n(?!\n)", " ", text)

def remove_headers_footers(text: str) -> str:
    return HEADER_FOOTER_CANDIDATE.sub("", text)

def clean_text(raw: str, config: Dict[str, Any] = None) -> str:
    cfg = {**DEFAULT_CONFIG, **(config or {})}
    s = raw

    if cfg["normalize_unicode"]:
        s = unicodedata.normalize("NFKC", s)
    if cfg["strip_zero_width"]:
        s = ZERO_WIDTH_PATTERN.sub("", s)
    if cfg["replace_nbsp"]:
        s = NBSP_PATTERN.sub(" ", s)
    if cfg["smart_quotes_to_ascii"]:
        s = smart_quotes_to_ascii(s)
    if cfg["fix_hyphenation"]:
        s = fix_hyphenation(s)
    if cfg["remove_bullets"]:
        s = BULLET_PATTERN.sub("", s)
    if cfg["remove_headers_footers"]:
        s = remove_headers_footers(s)
    if cfg["trim_lines"]:
        s = "\n".join(line.strip() for line in s.splitlines())
    if cfg["join_broken_lines"]:
        s = join_broken_lines(s)
    if cfg["deduplicate_spaces"]:
        s = re.sub(r"[ \t]{2,}", " ", s)
    if cfg["collapse_blank_lines"]:
        s = re.sub(r"\n{3,}", "\n\n", s).strip()
    if cfg["redact_emails"]:
        s = EMAIL_PATTERN.sub("[EMAIL_REDACTED]", s)
    if cfg["redact_phones"]:
        s = PHONE_PATTERN.sub("[PHONE_REDACTED]", s)
    if cfg["max_length"] and len(s) > cfg["max_length"]:
        cutoff = cfg["max_length"]
        m = re.search(r"[\.!\?]\s", s[cutoff-400:cutoff+200]) if cutoff > 400 else None
        if m:
            s = s[:cutoff-400 + m.end()].rstrip()
        else:
            s = s[:cutoff].rstrip() + "…"

    return s

def sha256_of_text(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8")).hexdigest()